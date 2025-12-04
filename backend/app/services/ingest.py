"""
Document Ingestion Service
Handles document processing, chunking, and embedding generation.
"""
import os
import hashlib
import uuid
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
import logging

import fitz  # PyMuPDF
from docx import Document as DocxDocument
import markdown
import tiktoken

from app.core.config import settings
from app.core.database import db
from app.models import Document, Chunk, DocumentStatus

logger = logging.getLogger(__name__)


class IngestService:
    """Service for document ingestion and processing."""
    
    def __init__(self):
        self.chunk_size = settings.RAG_CHUNK_SIZE
        self.chunk_overlap = settings.RAG_CHUNK_OVERLAP
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
    
    def save_uploaded_file(
        self,
        file,
        filename: str,
        user_id: str
    ) -> Tuple[Document, bool]:
        """
        Save an uploaded file and create a document record.
        Returns (document, is_duplicate)
        """
        # Generate unique filename
        ext = os.path.splitext(filename)[1].lower()
        unique_filename = f"{uuid.uuid4()}{ext}"
        file_path = os.path.join(settings.UPLOAD_PATH, unique_filename)
        
        # Ensure upload directory exists
        os.makedirs(settings.UPLOAD_PATH, exist_ok=True)
        
        # Save file
        file.save(file_path)
        
        # Compute file hash
        checksum = self._compute_file_hash(file_path)
        
        # Check for duplicate
        existing_doc = Document.query.filter_by(checksum=checksum).first()
        if existing_doc:
            # Remove the duplicate file
            os.remove(file_path)
            return existing_doc, True
        
        # Get file size
        file_size = os.path.getsize(file_path)
        
        # Determine file type
        file_type = ext.lstrip('.').lower()
        
        # Create document record
        document = Document(
            filename=unique_filename,
            original_filename=filename,
            file_path=file_path,
            file_type=file_type,
            file_size=file_size,
            checksum=checksum,
            status=DocumentStatus.PENDING,
            uploaded_by=user_id
        )
        
        db.session.add(document)
        db.session.commit()
        
        return document, False
    
    def process_document(self, document_id: str) -> Document:
        """
        Process a document: extract text, chunk, and prepare for embedding.
        """
        document = Document.query.get(document_id)
        if not document:
            raise ValueError(f"Document {document_id} not found")
        
        try:
            # Update status
            document.status = DocumentStatus.PROCESSING
            db.session.commit()
            
            # Extract text based on file type
            text_content, page_count = self._extract_text(document)
            
            # Update page count
            document.page_count = page_count
            
            # Chunk the text
            chunks = self._chunk_text(text_content, document.original_filename)
            
            # Store chunks in database
            for i, chunk_data in enumerate(chunks):
                chunk = Chunk(
                    document_id=document.id,
                    text=chunk_data['text'],
                    page_number=chunk_data.get('page'),
                    paragraph_number=chunk_data.get('paragraph'),
                    chunk_index=i,
                    token_count=chunk_data['token_count'],
                    chunk_metadata={
                        'document_name': document.original_filename,
                        'file_type': document.file_type
                    }
                )
                db.session.add(chunk)
            
            # Update document
            document.chunk_count = len(chunks)
            document.status = DocumentStatus.PROCESSED
            document.processed_at = datetime.utcnow()
            db.session.commit()
            
            logger.info(f"Document {document_id} processed: {len(chunks)} chunks")
            return document
            
        except Exception as e:
            document.status = DocumentStatus.FAILED
            document.error_message = str(e)
            db.session.commit()
            logger.error(f"Error processing document {document_id}: {e}")
            raise
    
    def _extract_text(self, document: Document) -> Tuple[List[Dict], int]:
        """
        Extract text from document based on file type.
        Returns list of {text, page, paragraph} dicts and page count.
        """
        file_type = document.file_type.lower()
        
        if file_type == 'pdf':
            return self._extract_pdf(document.file_path)
        elif file_type == 'docx':
            return self._extract_docx(document.file_path)
        elif file_type == 'md':
            return self._extract_markdown(document.file_path)
        elif file_type == 'txt':
            return self._extract_text_file(document.file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _extract_pdf(self, file_path: str) -> Tuple[List[Dict], int]:
        """Extract text from PDF file."""
        doc = fitz.open(file_path)
        content = []
        page_count = len(doc)  # Get page count before closing
        
        for page_num in range(page_count):
            page = doc.load_page(page_num)
            text = page.get_text("text")
            
            # Split by paragraphs
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
            
            for para_num, paragraph in enumerate(paragraphs):
                if paragraph:
                    content.append({
                        'text': paragraph,
                        'page': page_num + 1,
                        'paragraph': para_num + 1
                    })
        
        doc.close()
        return content, page_count
    
    def _extract_docx(self, file_path: str) -> Tuple[List[Dict], int]:
        """Extract text from DOCX file."""
        doc = DocxDocument(file_path)
        content = []
        
        for para_num, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                content.append({
                    'text': text,
                    'page': None,  # DOCX doesn't have page numbers
                    'paragraph': para_num + 1
                })
        
        # Estimate page count (rough estimate)
        total_text = ' '.join([c['text'] for c in content])
        page_count = max(1, len(total_text) // 3000)  # ~3000 chars per page
        
        return content, page_count
    
    def _extract_markdown(self, file_path: str) -> Tuple[List[Dict], int]:
        """Extract text from Markdown file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert markdown to plain text (remove formatting)
        html = markdown.markdown(md_content)
        # Simple HTML strip (for basic conversion)
        import re
        text = re.sub(r'<[^>]+>', '', html)
        
        # Split by sections/paragraphs
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        content = []
        for para_num, paragraph in enumerate(paragraphs):
            if paragraph:
                content.append({
                    'text': paragraph,
                    'page': None,
                    'paragraph': para_num + 1
                })
        
        return content, 1
    
    def _extract_text_file(self, file_path: str) -> Tuple[List[Dict], int]:
        """Extract text from plain text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Split by paragraphs
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        content = []
        for para_num, paragraph in enumerate(paragraphs):
            if paragraph:
                content.append({
                    'text': paragraph,
                    'page': None,
                    'paragraph': para_num + 1
                })
        
        return content, 1
    
    def _chunk_text(
        self,
        content: List[Dict],
        document_name: str
    ) -> List[Dict]:
        """
        Chunk text using recursive character splitting.
        Maintains page/paragraph references.
        """
        chunks = []
        current_chunk = []
        current_tokens = 0
        current_page = None
        current_paragraph = None
        
        for item in content:
            text = item['text']
            tokens = len(self.tokenizer.encode(text))
            
            # If single item is larger than chunk size, split it
            if tokens > self.chunk_size:
                # Flush current chunk first
                if current_chunk:
                    chunks.append({
                        'text': ' '.join(current_chunk),
                        'token_count': current_tokens,
                        'page': current_page,
                        'paragraph': current_paragraph
                    })
                    current_chunk = []
                    current_tokens = 0
                
                # Split large text
                split_chunks = self._split_large_text(text, item['page'], item['paragraph'])
                chunks.extend(split_chunks)
                continue
            
            # Check if adding this would exceed chunk size
            if current_tokens + tokens > self.chunk_size and current_chunk:
                # Save current chunk
                chunks.append({
                    'text': ' '.join(current_chunk),
                    'token_count': current_tokens,
                    'page': current_page,
                    'paragraph': current_paragraph
                })
                
                # Start new chunk with overlap
                overlap_text = current_chunk[-1] if current_chunk else ''
                overlap_tokens = len(self.tokenizer.encode(overlap_text))
                
                if overlap_tokens <= self.chunk_overlap:
                    current_chunk = [overlap_text] if overlap_text else []
                    current_tokens = overlap_tokens
                else:
                    current_chunk = []
                    current_tokens = 0
            
            # Add to current chunk
            current_chunk.append(text)
            current_tokens += tokens
            current_page = item['page']
            current_paragraph = item['paragraph']
        
        # Don't forget the last chunk
        if current_chunk:
            chunks.append({
                'text': ' '.join(current_chunk),
                'token_count': current_tokens,
                'page': current_page,
                'paragraph': current_paragraph
            })
        
        return chunks
    
    def _split_large_text(
        self,
        text: str,
        page: Optional[int],
        paragraph: Optional[int]
    ) -> List[Dict]:
        """Split text that's larger than chunk size."""
        chunks = []
        sentences = text.replace('!', '.').replace('?', '.').split('.')
        
        current_chunk = []
        current_tokens = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            tokens = len(self.tokenizer.encode(sentence))
            
            if current_tokens + tokens > self.chunk_size and current_chunk:
                chunks.append({
                    'text': '. '.join(current_chunk) + '.',
                    'token_count': current_tokens,
                    'page': page,
                    'paragraph': paragraph
                })
                current_chunk = []
                current_tokens = 0
            
            current_chunk.append(sentence)
            current_tokens += tokens
        
        if current_chunk:
            chunks.append({
                'text': '. '.join(current_chunk) + '.',
                'token_count': current_tokens,
                'page': page,
                'paragraph': paragraph
            })
        
        return chunks
    
    def _compute_file_hash(self, file_path: str) -> str:
        """Compute SHA256 hash of a file."""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    def delete_document(self, document_id: str) -> None:
        """Delete a document and its chunks."""
        document = Document.query.get(document_id)
        if not document:
            raise ValueError(f"Document {document_id} not found")
        
        # Delete file from disk
        if os.path.exists(document.file_path):
            os.remove(document.file_path)
        
        # Delete from database (chunks will cascade)
        db.session.delete(document)
        db.session.commit()
        
        logger.info(f"Document {document_id} deleted")
    
    def get_document(self, document_id: str) -> Optional[Document]:
        """Get a document by ID."""
        return Document.query.get(document_id)
    
    def get_documents(
        self,
        page: int = 1,
        per_page: int = 20,
        status: Optional[DocumentStatus] = None,
        user_id: Optional[str] = None
    ) -> Tuple[List[Document], int]:
        """Get paginated documents, optionally filtered by user."""
        query = Document.query
        
        # Filter by user if specified
        if user_id:
            query = query.filter_by(uploaded_by=user_id)
        
        if status:
            query = query.filter_by(status=status)
        
        query = query.order_by(Document.created_at.desc())
        
        total = query.count()
        documents = query.offset((page - 1) * per_page).limit(per_page).all()
        
        return documents, total
